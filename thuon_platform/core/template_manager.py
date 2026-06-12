# core/template_manager.py

import yaml
import logging
from pathlib import Path
from jinja2 import Environment, BaseLoader

logger = logging.getLogger('thuon.template_manager')

_DEFAULT_TEMPLATE_PATH = Path(__file__).parent.parent / 'data' / 'templates.yaml'


class TemplateManager:
	def __init__(self, template_file_path: str = str(_DEFAULT_TEMPLATE_PATH)):
		self.template_file_path = template_file_path
		self.templates: dict = {}
		self.load_templates()
		self._jinja_env = Environment(loader=BaseLoader())

	def load_templates(self) -> dict:
		try:
			with open(self.template_file_path, 'r') as f:
				self.templates = yaml.safe_load(f) or {}
		except FileNotFoundError:
			logger.warning(f"Template file not found: {self.template_file_path}")
			self.templates = {}
		return self.templates

	def get_template(self, template_name: str) -> dict | None:
		return self.templates.get(template_name)

	def list_templates(self, template_type: str | None = None) -> list[str]:
		if template_type is None:
			return list(self.templates.keys())
		return [k for k, v in self.templates.items() if isinstance(v, dict) and v.get('type') == template_type]

	def validate_template(self, template_dict: dict, template_type: str) -> bool:
		required_keys = {'content', 'type'}
		return required_keys.issubset(template_dict.keys()) and template_dict.get('type') == template_type

	def apply_data_to_template(self, template_dict: dict, data_dict: dict) -> str:
		content = template_dict.get('content', '') if isinstance(template_dict, dict) else str(template_dict)
		try:
			tmpl = self._jinja_env.from_string(content)
			return tmpl.render(**data_dict)
		except Exception as e:
			logger.error(f"Template render error: {e}")
			return content

	def generate_document_from_template_string(self, template_string: str, data_dict: dict, output_format: str, output_path: str) -> bool:
		try:
			rendered = self._jinja_env.from_string(template_string).render(**data_dict)
			Path(output_path).parent.mkdir(parents=True, exist_ok=True)

			if output_format == 'docx':
				from docx import Document
				doc = Document()
				for line in rendered.split('\n'):
					line = line.strip()
					if line.startswith('# '):
						doc.add_heading(line[2:], level=1)
					elif line.startswith('## '):
						doc.add_heading(line[3:], level=2)
					elif line.startswith('### '):
						doc.add_heading(line[4:], level=3)
					elif line:
						doc.add_paragraph(line)
				doc.save(output_path)
			else:
				with open(output_path, 'w', encoding='utf-8') as f:
					f.write(rendered)
			return True
		except Exception as e:
			logger.error(f"Document generation error: {e}")
			return False
